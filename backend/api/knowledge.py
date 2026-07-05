"""
知识库文档上传/管理 API
支持 PDF、DOCX、MD、TXT、PNG、JPG、JPEG 的上传、分块、Embedding、入库
"""
import os
import logging
logger = logging.getLogger(__name__)
import uuid
from typing import List, Optional

from fastapi import APIRouter, UploadFile, File, HTTPException
from sqlalchemy import select, delete, func
from sqlalchemy.ext.asyncio import AsyncSession

from database import AsyncSessionLocal
from models.knowledge import KnowledgeDocument, MedicalKnowledgeVector

router = APIRouter(tags=["知识库"])

KNOWLEDGE_UPLOAD_DIR = "uploads/knowledge"


def _ensure_dir(path: str) -> None:
    if not os.path.exists(path):
        os.makedirs(path, exist_ok=True)


async def _extract_text_from_image(image_path: str) -> str:
    logger.info(f"[Knowledge API] 图片OCR提取: {image_path}")
    try:
        from api.upload import extract_text_from_image as _up_extract
        from pathlib import Path
        return await _up_extract(Path(image_path))
    except Exception as e:
        logger.error(f"[Knowledge API] 图片OCR失败: {e}")
        return ""


async def _extract_text_from_pdf(pdf_path: str) -> str:
    logger.info(f"[Knowledge API] PDF文本提取: {pdf_path}")
    try:
        from api.upload import extract_text_from_pdf as _up_pdf
        from pathlib import Path
        return _up_pdf(Path(pdf_path)) or ""
    except Exception as e:
        logger.error(f"[Knowledge API] PDF提取失败: {e}")
        return ""


async def _extract_text_from_docx(docx_path: str) -> str:
    logger.info(f"[Knowledge API] DOCX文本提取: {docx_path}")
    try:
        import docx  # python-docx
        doc = docx.Document(docx_path)
        parts = []
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                parts.append(para.text)
        return "\n".join(parts)
    except ImportError:
        logger.info("[Knowledge API] python-docx 未安装，尝试备用方式")
        try:
            import subprocess
            result = subprocess.run(
                ["pandoc", "-t", "plain", docx_path],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0:
                return result.stdout
            logger.error(f"[Knowledge API] pandoc 失败: {result.stderr}")
        except Exception as e:
            logger.error(f"[Knowledge API] pandoc 也失败: {e}")
        return ""
    except Exception as e:
        logger.error(f"[Knowledge API] DOCX解析异常: {e}")
        return ""


async def _generate_embedding(text: str, settings) -> Optional[List[float]]:
    """调用 DashScope 生成 embedding"""
    if not text or not text.strip():
        logger.info("[Knowledge API] 空文本，跳过Embedding")
        return None
    try:
        import dashscope
        resp = dashscope.TextEmbedding.call(
            model=settings.embedding_model,
            api_key=settings.openai_api_key,
            input=text,
        )
        if resp.status_code != 200:
            logger.error(f"[Knowledge API] Embedding失败: code={resp.code}, msg={resp.message}")
            return None
        vector = resp.output["embeddings"][0]["embedding"]
        logger.info(f"[Knowledge API] Embedding成功，维度: {len(vector)}")
        return vector
    except Exception as e:
        logger.error(f"[Knowledge API] Embedding异常: {e}")
        return None


async def _insert_vector(
    session: AsyncSession,
    document_id: str,
    chunk_index: int,
    content: str,
    embedding: List[float],
    source: str = "user_upload",
) -> None:
    """使用 CAST 函数代替 :: 转型，避免 SQLAlchemy 参数解析冲突"""
    import json as _json
    from sqlalchemy import text

    vector_str = "[" + ",".join(str(x) for x in embedding) + "]"
    metadata_json = _json.dumps({
        "document_id": document_id,
        "chunk_index": chunk_index,
    }, ensure_ascii=False)

    stmt = text("""
        INSERT INTO medical_knowledge_vectors (id, document_id, chunk_index, content, embedding, source, extra_metadata)
        VALUES (
            CAST(:id AS uuid),
            CAST(:doc_id AS uuid),
            :chunk_index,
            :content,
            CAST(:embedding AS vector),
            :source,
            :extra_metadata
        )
    """)

    await session.execute(stmt, {
        "id": str(uuid.uuid4()),
        "doc_id": document_id,
        "chunk_index": chunk_index,
        "content": content,
        "embedding": vector_str,
        "source": source,
        "extra_metadata": metadata_json,
    })


_ALLOWED_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/markdown",
    "text/plain",
    "image/png",
    "image/jpeg",
    "image/jpg",
}


@router.post("/api/knowledge/upload")
async def upload_knowledge(file: UploadFile = File(...)):
    logger.info(f"[Knowledge API] 收到上传请求: {file.filename}, 类型: {file.content_type}")

    if file.content_type not in _ALLOWED_TYPES:
        logger.info(f"[Knowledge API] 不支持的文件类型: {file.content_type}")
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: {file.content_type}")

    _ensure_dir(KNOWLEDGE_UPLOAD_DIR)
    safe_name = f"{uuid.uuid4()}_{file.filename}"
    file_path = os.path.join(KNOWLEDGE_UPLOAD_DIR, safe_name)

    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)
    file_size = len(content)
    logger.info(f"[Knowledge API] 文件已保存: {file_path}, 大小: {file_size} 字节")

    extracted_text = ""
    try:
        if file.content_type and file.content_type.startswith("image/"):
            extracted_text = await _extract_text_from_image(file_path)
        elif file.content_type == "application/pdf":
            extracted_text = await _extract_text_from_pdf(file_path)
        elif file.content_type in ("text/plain", "text/markdown"):
            with open(file_path, "r", encoding="utf-8") as f:
                extracted_text = f.read()
        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            extracted_text = await _extract_text_from_docx(file_path)
        else:
            logger.info(f"[Knowledge API] 未匹配的内容类型: {file.content_type}")
    except Exception as e:
        logger.error(f"[Knowledge API] 文本提取阶段异常: {e}")

    logger.info(f"[Knowledge API] 文本提取完成: {len(extracted_text)} 字符")
    if not extracted_text or not extracted_text.strip():
        raise HTTPException(status_code=400, detail="无法从文件中提取文本")

    from langchain_text_splitters import RecursiveCharacterTextSplitter
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = splitter.split_text(extracted_text)
    logger.info(f"[Knowledge API] 文本分块完成: {len(chunks)} 块")

    from config import get_settings
    settings = get_settings()
    logger.info(f"[Knowledge API] 开始生成 Embedding... 共 {len(chunks)} 块")

    doc_id = str(uuid.uuid4())
    success_chunks = 0

    async with AsyncSessionLocal() as session:
        for i, chunk in enumerate(chunks):
            embedding = await _generate_embedding(chunk, settings)
            if embedding:
                try:
                    await _insert_vector(
                        session=session,
                        document_id=doc_id,
                        chunk_index=i,
                        content=chunk,
                        embedding=embedding,
                        source="user_upload",
                    )
                    success_chunks += 1
                    if (i + 1) % 5 == 0 or i == len(chunks) - 1:
                        logger.info(f"[Knowledge API] Embedding进度: {i + 1}/{len(chunks)}")
                except Exception as e:
                    logger.error(f"[Knowledge API] 写入向量异常 块{i+1}: {e}")
            else:
                logger.error(f"[Knowledge API] Embedding失败: 块 {i + 1}")

        await session.commit()
        logger.info(f"[Knowledge API] 向量已写入数据库: {success_chunks} 条记录")

        doc = KnowledgeDocument(
            id=doc_id,
            title=file.filename,
            filename=file.filename,
            source="用户上传",
            file_type=file.content_type or "",
            file_path=file_path,
            file_size=file_size,
            chunk_count=success_chunks,
            status="active",
        )
        session.add(doc)
        await session.commit()

    logger.info(f"[Knowledge API] 文档记录已创建: doc_id={doc_id}, title={file.filename}")
    return {
        "success": True,
        "doc_id": doc_id,
        "chunks": success_chunks,
        "title": file.filename,
    }


@router.get("/api/knowledge/documents")
async def list_documents(page: int = 1, page_size: int = 20):
    logger.info(f"[Knowledge API] 查询文档列表: page={page}, size={page_size}")

    async with AsyncSessionLocal() as session:
        count_result = await session.execute(select(func.count(KnowledgeDocument.id)))
        total = count_result.scalar() or 0

        result = await session.execute(
            select(KnowledgeDocument)
            .order_by(KnowledgeDocument.created_at.desc())
            .offset((page - 1) * page_size)
            .limit(page_size)
        )
        docs = result.scalars().all()

        data = []
        for d in docs:
            data.append({
                "id": d.id,
                "title": getattr(d, "title", None) or d.filename,
                "filename": d.filename,
                "source": getattr(d, "source", None) or "user_upload",
                "file_type": d.file_type,
                "file_size": d.file_size,
                "chunk_count": d.chunk_count,
                "status": d.status,
                "created_at": d.created_at.isoformat() if d.created_at else None,
            })

    logger.info(f"[Knowledge API] 查询完成: 总数={total}, 返回={len(data)}")
    return {
        "success": True,
        "data": data,
        "total": total,
        "page": page,
        "page_size": page_size,
    }


@router.delete("/api/knowledge/documents/{doc_id}")
async def delete_document(doc_id: str):
    logger.info(f"[Knowledge API] 删除文档请求: doc_id={doc_id}")

    async with AsyncSessionLocal() as session:
        result = await session.execute(
            select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id)
        )
        doc = result.scalar_one_or_none()
        if not doc:
            raise HTTPException(status_code=404, detail="文档不存在")

        delete_result = await session.execute(
            delete(MedicalKnowledgeVector).where(MedicalKnowledgeVector.document_id == doc_id)
        )
        deleted_vectors = delete_result.rowcount or 0
        logger.info(f"[Knowledge API] 向量已删除: {deleted_vectors} 条")

        file_path = doc.file_path
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info(f"[Knowledge API] 物理文件已删除: {file_path}")
            except Exception as e:
                logger.error(f"[Knowledge API] 物理文件删除失败: {e}")

        await session.delete(doc)
        await session.commit()

    logger.info(f"[Knowledge API] 文档已删除: doc_id={doc_id}")
    return {
        "success": True,
        "doc_id": doc_id,
        "deleted_vectors": deleted_vectors,
    }
