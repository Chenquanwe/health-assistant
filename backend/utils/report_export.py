import io
import logging
logger = logging.getLogger(__name__)
import os
import re
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


def _parse_markdown(text: str):
    elements = []
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        # ----- 识别 **数字. 标题** 格式（LLM 生成的报告常用） -----
        bold_title_match = re.match(r'^\*\*(\d+)\.\s*(.*?)\*\*$', line)
        if bold_title_match:
            title_text = f"{bold_title_match.group(1)}. {bold_title_match.group(2)}"
            elements.append(('h2', title_text))
            i += 1
            continue

        # ----- 表格检测 -----
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = [line]
            i += 2
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].rstrip())
                i += 1
            elements.append(('table', table_lines))
            continue

        # ----- 标准 Markdown 标题 -----
        if line.startswith('# '):
            elements.append(('h1', line[2:].strip()))
        elif line.startswith('## '):
            elements.append(('h2', line[3:].strip()))
        elif line.startswith('### '):
            elements.append(('h3', line[4:].strip()))
        elif line.startswith('> '):
            elements.append(('quote', line[2:].strip()))
        elif line.startswith('- '):
            elements.append(('list_item', line[2:].strip()))
        else:
            elements.append(('paragraph', line.strip()))
        i += 1
    return elements

def _split_bold_parts(text: str):
    parts = []
    pattern = re.compile(r'\*\*(.*?)\*\*')
    last_end = 0
    for match in pattern.finditer(text):
        if match.start() > last_end:
            parts.append((text[last_end:match.start()], False))
        parts.append((match.group(1), True))
        last_end = match.end()
    if last_end < len(text):
        parts.append((text[last_end:], False))
    if not parts:
        parts.append((text, False))
    return parts


def _is_emoji(c):
    """判断字符是否为需要特殊处理的 emoji（辅助平面字符）"""
    return ord(c) > 0xFFFF


def _add_run_with_emoji(paragraph, text, font_name, font_size, is_bold, color=None):
    """添加 run，自动将 emoji 字符用 Segoe UI Emoji 字体输出"""
    if not text:
        return

    buf = []
    for ch in text:
        if _is_emoji(ch):
            if buf:
                run = paragraph.add_run(''.join(buf))
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                run.font.size = font_size
                run.bold = is_bold
                if color:
                    run.font.color.rgb = color
                buf.clear()
            run = paragraph.add_run(ch)
            run.font.name = 'Segoe UI Emoji'
            run.font.size = font_size
            run.bold = is_bold
        else:
            buf.append(ch)

    if buf:
        run = paragraph.add_run(''.join(buf))
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = font_size
        run.bold = is_bold
        if color:
            run.font.color.rgb = color


# --------------------------- PDF ---------------------------

def markdown_to_pdf(md_text: str) -> bytes:
    pdf = FPDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=15)

    font_dir = os.path.join(os.path.dirname(__file__), 'fonts')
    hei_path = os.path.join(font_dir, 'SimHei.ttf')
    sun_path = os.path.join(font_dir, 'SimSun.ttf')
    has_chinese = os.path.exists(hei_path) and os.path.exists(sun_path)

    if has_chinese:
        pdf.add_font('SimHei', '', hei_path, uni=True)
        pdf.add_font('SimSun', '', sun_path, uni=True)
        title_font = 'SimHei'
        body_font = 'SimSun'
    else:
        title_font = 'Helvetica'
        body_font = 'Helvetica'

    pdf.add_page()
    elements = _parse_markdown(md_text)

    # emoji → 文本映射
    emoji_map = {
        '🟢': '[绿色]',
        '🔴': '[红色]',
        '🟡': '[黄色]',
        '⚠️': '[警告]',
        '✅': '[完成]',
        '❌': '[错误]',
    }

    for elem_type, content in elements:
        if isinstance(content, str):
            # 去除 ** 标记
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
            # 替换 emoji
            for emoji, text in emoji_map.items():
                content = content.replace(emoji, text)

        if elem_type == 'h1':
            pdf.set_font(title_font, '', 16)
            pdf.multi_cell(pdf.epw, 7, content)
            pdf.ln(3)
        elif elem_type == 'h2':
            pdf.set_font(title_font, '', 14)
            pdf.multi_cell(pdf.epw, 6, content)
            pdf.ln(2)
        elif elem_type == 'h3':
            pdf.set_font(title_font, '', 12)
            pdf.multi_cell(pdf.epw, 6, content)
            pdf.ln(2)
        elif elem_type == 'list_item':
            pdf.set_font(body_font, '', 10)
            list_indent = 10  # 列表整体左缩进 10mm（约两个汉字）
            pdf.cell(list_indent, 6, '')  # 缩进空白
            pdf.cell(8, 6, '•')  # 圆点
            pdf.multi_cell(pdf.epw - list_indent - 8, 6, content)
            pdf.ln(1)
        elif elem_type == 'quote':
            pdf.set_font(body_font, '', 10)
            pdf.set_x(pdf.l_margin + 10)
            pdf.multi_cell(pdf.epw - 10, 6, content)
            pdf.ln(2)
        elif elem_type == 'table':
            continue
        else:
            pdf.set_font(body_font, '', 10)
            pdf.multi_cell(pdf.epw, 6, '　　' + content)  # 全角空格实现首行缩进两个汉字
            pdf.ln(1)

    return bytes(pdf.output())


# --------------------------- Word ---------------------------

def markdown_to_docx(md_text: str) -> bytes:
    doc = Document()

    # ===================== 页面设置 =====================
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ===================== 正文样式 =====================
    normal_style = doc.styles['Normal']
    normal_style.font.name = '宋体'
    normal_style.font.size = Pt(14)                      # 四号
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal_style.paragraph_format.line_spacing = Pt(20)  # 固定值20磅
    normal_style.paragraph_format.first_line_indent = Pt(28)   # 首行缩进2字符
    normal_style.paragraph_format.space_before = Pt(6)         # 0.5行
    normal_style.paragraph_format.space_after = Pt(6)

    # 不再修改内置标题样式，标题将使用普通段落手动设置大纲级别

    elements = _parse_markdown(md_text)

    for elem_type, content in elements:
        logger.info(f"[DEBUG] 元素类型: {elem_type}, 内容: {str(content)[:20]}")  # 加这行
        # ----- 标题（使用普通段落+大纲级别，完全自定义格式） -----
        if elem_type in ('h1', 'h2', 'h3'):
            level = {'h1': 1, 'h2': 2, 'h3': 3}[elem_type]

            p = doc.add_paragraph()                    # 普通段落，不继承任何标题样式
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 设置大纲级别，让标题在导航窗格中显示
            pf = p.paragraph_format
            pf.outline_level = level - 1               # 0-based (0 = 一级标题, 1 = 二级...)
            pf.first_line_indent = Pt(0)               # 无首行缩进
            pf.left_indent = Pt(0)                     # 无左缩进
            pf.right_indent = Pt(0)                    # 无右缩进
            pf.space_before = Pt(6)                    # 段前 0.5 行
            pf.space_after = Pt(6)                     # 段后 0.5 行
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pf.line_spacing = 1.0                      # 单倍行距

            # 字体格式
            run = p.add_run(content)
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(16)                     # 三号
            run.font.bold = True
            if elem_type == 'h1':
                run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)  # 一级标题蓝色
            else:
                run.font.color.rgb = RGBColor(0, 0, 0)

        # ----- 普通段落 -----
        elif elem_type == 'paragraph':
            p = doc.add_paragraph()
            parts = _split_bold_parts(content)
            for text, is_bold in parts:
                _add_run_with_emoji(p, text, '宋体', Pt(14), is_bold)

        # ----- 列表项 -----
        elif elem_type == 'list_item':
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.75)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(6)
            parts = _split_bold_parts(content)
            for text, is_bold in parts:
                _add_run_with_emoji(p, text, '宋体', Pt(11), is_bold)

        # ----- 引用 / 免责声明 -----
        elif elem_type == 'quote':
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Inches(0.5)
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(20)
            pf.first_line_indent = Pt(0)

            if content.startswith('⚠️'):
                run_symbol = p.add_run('⚠️ ')
                run_symbol.font.name = 'Segoe UI Emoji'
                run_symbol.font.size = Pt(14)
                run_symbol.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)
                run_symbol.bold = False

                remaining = content[2:].lstrip()
                run_text = p.add_run(remaining)
                run_text.font.name = '黑体'
                run_text._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run_text.font.size = Pt(14)
                run_text.bold = True
            else:
                parts = _split_bold_parts(content)
                for text, is_bold in parts:
                    _add_run_with_emoji(p, text, '宋体', Pt(14), is_bold,
                                        color=RGBColor(0x66, 0x66, 0x66))

        # ----- 表格 -----
        elif elem_type == 'table':
            table_data = []
            for row_line in content:
                if '---' in row_line:
                    continue
                cells = [c.strip() for c in row_line.split('|') if c.strip()]
                if cells:
                    table_data.append(cells)
            if not table_data:
                continue

            num_cols = len(table_data[0])
            num_rows = len(table_data)
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            table.autofit = True

            for row_idx, row in enumerate(table_data):
                cells = table.rows[row_idx].cells
                for col_idx, cell in enumerate(row):
                    if col_idx < len(cells):
                        paragraph = cells[col_idx].paragraphs[0]
                        run = paragraph.add_run(cell)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(10)
                        if row_idx == 0:
                            run.bold = True
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()